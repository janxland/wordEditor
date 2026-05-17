"""快速校验 hutb-carbon-neutral 输出 docx 的样式/页眉页脚/引用/数学。"""
import sys
from collections import Counter
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

OUT = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("output/carbon-neutral-renewable-hutb-carbon-neutral.docx")
d = Document(str(OUT))
print(f"文件: {OUT}  大小: {OUT.stat().st_size} bytes")
print("段落数:", len(d.paragraphs))

styles = Counter(p.style.name for p in d.paragraphs)
print("样式分布:")
for k, v in sorted(styles.items(), key=lambda x: -x[1]):
    print(f"  {k:20s} {v}")

sec = d.sections[0]
hdr_text = "\n".join(p.text for p in sec.header.paragraphs if p.text.strip())
print("页眉文本:", repr(hdr_text[:120]))

ftr_xml = sec.footer._element.xml
print("页脚含 PAGE 域:", " PAGE " in ftr_xml or ">PAGE<" in ftr_xml)
print("页脚含 NUMPAGES 域:", "NUMPAGES" in ftr_xml)
print("页脚 fldChar 数:", ftr_xml.count("fldChar"))
ftr_text = "".join(p.text for p in sec.footer.paragraphs)
print("页脚可见文本:", repr(ftr_text))

body_xml = d.element.xml
ref_count = body_xml.count(" REF Ref")
sup_tag = '<w:vertAlign w:val="superscript"/>'
sup_count = body_xml.count(sup_tag)
print("正文 REF Ref 域数:", ref_count)
print("上标 run 数:", sup_count)
bms = [b for b in d.element.iter(qn("w:bookmarkStart")) if (b.get(qn("w:name")) or "").startswith("Ref")]
print("Ref* 书签数:", len(bms))

import re
math_re = re.compile(r"[ηωσΔΣ×≤≥₀₁₂₃ᵢₙₛₑₜₕᵣ⁰⁻⁽⁾·₍₎]")
hits = [p.text for p in d.paragraphs if math_re.search(p.text)]
print("含 Unicode 数学段落数:", len(hits))
for h in hits[:3]:
    print("  >", h[:140])

# 标题层级统计
hcounts = Counter()
for p in d.paragraphs:
    name = p.style.name
    if name.startswith("Heading"):
        hcounts[name] += 1
print("Heading 分布:", dict(hcounts))
